"""
多模态文档加载器

扩展DocumentLoader以支持PDF、DOCX等格式，
并提供图片理解能力（通过多模态LLM）。
"""
import os
import logging
from typing import List, Optional, Dict, Any
from dataclasses import dataclass

from src.core.document_loader import Document, DocumentLoader

logger = logging.getLogger(__name__)


@dataclass
class MultimodalDocument(Document):
    """多模态文档"""
    media_type: Optional[str] = None
    image_description: Optional[str] = None


class MultimodalDocumentLoader(DocumentLoader):
    """多模态文档加载器
    
    在原有DocumentLoader基础上扩展：
    - PDF文档解析（使用pypdf）
    - DOCX文档解析（使用python-docx）
    - 图片描述生成（通过多模态LLM）
    """
    
    SUPPORTED_EXTENSIONS = {
        '.txt', '.md', '.json',
        '.pdf', '.docx', '.doc',
        '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'
    }
    
    def load_file(self, file_path: str) -> Optional[MultimodalDocument]:
        """加载单个文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            多模态文档对象
        """
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            return None
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext in ('.txt', '.md', '.json'):
            doc = super().load_file(file_path)
            if doc:
                return MultimodalDocument(
                    doc_id=doc.doc_id,
                    content=doc.content,
                    metadata=doc.metadata,
                    source=doc.source,
                    doc_type=doc.doc_type
                )
            return None
        
        elif ext == '.pdf':
            return self._load_pdf(file_path)
        
        elif ext in ('.docx', '.doc'):
            return self._load_docx(file_path)
        
        elif ext in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'):
            return self._load_image(file_path)
        
        else:
            logger.warning(f"不支持的文件格式: {ext}")
            return None
    
    def _load_pdf(self, file_path: str) -> Optional[MultimodalDocument]:
        """加载PDF文档"""
        try:
            import pypdf
            
            reader = pypdf.PdfReader(file_path)
            text_parts = []
            
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(f"[第{page_num + 1}页]\n{text.strip()}")
            
            content = "\n\n".join(text_parts)
            
            if not content.strip():
                logger.warning(f"PDF文档无文本内容: {file_path}")
                return None
            
            return MultimodalDocument(
                doc_id=self._generate_doc_id(file_path),
                content=content,
                metadata={
                    "source": file_path,
                    "format": "pdf",
                    "page_count": len(reader.pages),
                    "doc_type": "pdf"
                },
                source=file_path,
                doc_type="pdf",
                media_type="application/pdf"
            )
            
        except ImportError:
            logger.error("pypdf未安装，无法解析PDF。请运行: pip install pypdf")
            return None
        except Exception as e:
            logger.error(f"PDF文档加载失败 {file_path}: {e}")
            return None
    
    def _load_docx(self, file_path: str) -> Optional[MultimodalDocument]:
        """加载DOCX文档"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            content = "\n\n".join(text_parts)
            
            if not content.strip():
                logger.warning(f"DOCX文档无文本内容: {file_path}")
                return None
            
            return MultimodalDocument(
                doc_id=self._generate_doc_id(file_path),
                content=content,
                metadata={
                    "source": file_path,
                    "format": "docx",
                    "doc_type": "docx"
                },
                source=file_path,
                doc_type="docx",
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        except ImportError:
            logger.error("python-docx未安装，无法解析DOCX。请运行: pip install python-docx")
            return None
        except Exception as e:
            logger.error(f"DOCX文档加载失败 {file_path}: {e}")
            return None
    
    def _load_image(self, file_path: str) -> Optional[MultimodalDocument]:
        """加载图片（生成描述）"""
        ext = os.path.splitext(file_path)[1].lower()
        mime_map = {
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.gif': 'image/gif',
            '.bmp': 'image/bmp',
            '.webp': 'image/webp'
        }
        
        description = self._describe_image(file_path)
        
        content = description or f"[图片文件: {os.path.basename(file_path)}]"
        
        return MultimodalDocument(
            doc_id=self._generate_doc_id(file_path),
            content=content,
            metadata={
                "source": file_path,
                "format": ext.lstrip('.'),
                "doc_type": "image",
                "has_description": description is not None
            },
            source=file_path,
            doc_type="image",
            media_type=mime_map.get(ext, "image/unknown"),
            image_description=description
        )
    
    def _describe_image(self, file_path: str) -> Optional[str]:
        """使用多模态LLM生成图片描述"""
        try:
            from src.core.llm import get_llm
            import base64
            
            with open(file_path, "rb") as f:
                image_data = base64.b64encode(f.read()).decode()
            
            ext = os.path.splitext(file_path)[1].lower()
            mime_map = {
                '.png': 'image/png',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
            }
            mime_type = mime_map.get(ext, 'image/jpeg')
            
            llm = get_llm()
            
            if hasattr(llm, 'invoke'):
                from langchain_core.messages import HumanMessage
                
                message = HumanMessage(content=[
                    {"type": "text", "text": "请描述这张与宠物健康相关的图片内容，包括任何可见的症状、体征或医疗信息。"},
                    {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{image_data}"}}
                ])
                
                response = llm.invoke([message])
                return response.content if hasattr(response, 'content') else str(response)
            
            return None
            
        except ImportError:
            logger.debug("多模态LLM不可用，跳过图片描述")
            return None
        except Exception as e:
            logger.debug(f"图片描述生成失败: {e}")
            return None
    
    def _generate_doc_id(self, file_path: str) -> str:
        """生成文档ID"""
        import hashlib
        return hashlib.md5(file_path.encode()).hexdigest()[:12]
